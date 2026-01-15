"""
Python script to generate Word documents from Dart files with explanations
"""

import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing required package: python-docx")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# File explanations mapping - Vietnamese
file_explanations = {
    "main.dart": """MỤC ĐÍCH FILE:
Đây là điểm khởi đầu của ứng dụng Flutter. File này khởi tạo ứng dụng, thiết lập xác thực Supabase, cấu hình đa ngôn ngữ và quản lý luồng xác thực người dùng.

CÁC CHỨC NĂNG CHÍNH:
- main(): Khởi tạo Flutter bindings và Supabase, sau đó chạy ứng dụng
- MedEquipApp: Widget gốc cấu hình MaterialApp với theme, đa ngôn ngữ và định tuyến
- AuthGate: Xử lý kiểm tra trạng thái xác thực và chuyển hướng người dùng đến các màn hình phù hợp dựa trên trạng thái đăng nhập""",

    # Constants
    "constants/app_colors.dart": """MỤC ĐÍCH FILE:
Định nghĩa bảng màu được sử dụng trong toàn bộ ứng dụng để đảm bảo giao diện nhất quán.

CÁC CHỨC NĂNG CHÍNH:
- Cung cấp các định nghĩa màu tập trung
- Đảm bảo thiết kế trực quan nhất quán trong toàn ứng dụng""",

    "constants/app_theme.dart": """MỤC ĐÍCH FILE:
Định nghĩa cấu hình theme của ứng dụng bao gồm cài đặt theme sáng và tối.

CÁC CHỨC NĂNG CHÍNH:
- Cấu hình các thuộc tính theme của MaterialApp
- Thiết lập các bảng màu, kiểu chữ và theme cho các component
- Cung cấp giao diện nhất quán cho toàn bộ ứng dụng""",

    "constants/constants.dart": """MỤC ĐÍCH FILE:
Chứa các hằng số toàn cục được sử dụng trong toàn ứng dụng như API keys, URLs và các giá trị cấu hình.

CÁC CHỨC NĂNG CHÍNH:
- Lưu trữ Supabase URL và API keys
- Định nghĩa các giá trị hằng số của ứng dụng
- Tập trung quản lý cấu hình""",

    "constants/database_translations.dart": """MỤC ĐÍCH FILE:
Quản lý bản dịch và ánh xạ các giá trị cơ sở dữ liệu sang văn bản thân thiện với người dùng bằng nhiều ngôn ngữ.

CÁC CHỨC NĂNG CHÍNH:
- Dịch các giá trị enum của database sang văn bản dễ hiểu
- Hỗ trợ đa ngôn ngữ
- Ánh xạ các mã trạng thái sang nhãn dễ đọc""",

    # Models
    "models/audit_log.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu cho các bản ghi nhật ký kiểm toán theo dõi tất cả hoạt động và thay đổi trong hệ thống.

CÁC CHỨC NĂNG CHÍNH:
- Đại diện cho các bản ghi theo dõi hoạt động
- Theo dõi ai, làm gì, khi nào và ở đâu
- Cung cấp serialization/deserialization cho các thao tác database""",

    "models/borrow_request.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu đại diện cho các yêu cầu mượn thiết bị do người dùng tạo.

CÁC CHỨC NĂNG CHÍNH:
- Lưu trữ thông tin chi tiết về việc mượn (người mượn, thiết bị, ngày tháng, trạng thái)
- Quản lý vòng đời yêu cầu mượn
- Cung cấp serialization/deserialization cho các thao tác database""",

    "models/category.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu cho các danh mục thiết bị được sử dụng để tổ chức và phân loại thiết bị y tế.

CÁC CHỨC NĂNG CHÍNH:
- Đại diện cho các danh mục thiết bị
- Lưu trữ metadata của danh mục (tên, mô tả, ID)
- Cung cấp serialization/deserialization cho các thao tác database""",

    "models/equipment.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu cốt lõi đại diện cho các mục thiết bị y tế trong hệ thống.

CÁC CHỨC NĂNG CHÍNH:
- Lưu trữ thông tin thiết bị toàn diện (tên, số serial, trạng thái, danh mục, vị trí)
- Quản lý vòng đời và tình trạng sẵn có của thiết bị
- Cung cấp serialization/deserialization cho các thao tác database""",

    "models/inventory_log.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu để theo dõi các thay đổi hàng tồn kho và chuyển động kho.

CÁC CHỨC NĂNG CHÍNH:
- Ghi lại các giao dịch hàng tồn kho
- Theo dõi thay đổi số lượng và lý do
- Cung cấp nhật ký kiểm toán cho quản lý hàng tồn kho""",

    "models/user.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu đại diện cho người dùng hệ thống với vai trò và quyền hạn của họ.

CÁC CHỨC NĂNG CHÍNH:
- Lưu trữ thông tin hồ sơ người dùng
- Quản lý vai trò và dữ liệu xác thực của người dùng
- Cung cấp serialization/deserialization cho các thao tác database""",

    "models/user_settings.dart": """MỤC ĐÍCH FILE:
Mô hình dữ liệu để lưu trữ tùy chọn người dùng và cài đặt ứng dụng.

CÁC CHỨC NĂNG CHÍNH:
- Quản lý cài đặt cụ thể của người dùng (ngôn ngữ, thông báo, theme)
- Lưu trữ tùy chọn người dùng
- Cung cấp serialization/deserialization cho các thao tác database""",

    # Providers
    "providers/locale_provider.dart": """MỤC ĐÍCH FILE:
Provider quản lý trạng thái để xử lý thay đổi ngôn ngữ/locale của ứng dụng.

CÁC CHỨC NĂNG CHÍNH:
- Quản lý trạng thái locale hiện tại
- Lưu trữ tùy chọn ngôn ngữ
- Thông báo cho các widget khi locale thay đổi
- Tích hợp với cài đặt người dùng""",

    # Services
    "services/audit_log_service.dart": """MỤC ĐÍCH FILE:
Lớp service để quản lý nhật ký kiểm toán - ghi lại và truy xuất nhật ký hoạt động của hệ thống.

CÁC CHỨC NĂNG CHÍNH:
- Tạo các mục nhật ký kiểm toán cho tất cả các hành động hệ thống
- Truy xuất nhật ký kiểm toán với bộ lọc và phân trang
- Cung cấp dấu vết kiểm toán để tuân thủ và gỡ lỗi""",

    "services/auth_service.dart": """MỤC ĐÍCH FILE:
Xử lý tất cả các thao tác xác thực bao gồm đăng nhập, đăng ký, đặt lại mật khẩu và quản lý phiên.

CÁC CHỨC NĂNG CHÍNH:
- Quản lý xác thực người dùng với Supabase
- Xử lý các thao tác đăng nhập/đăng ký/đăng xuất
- Quản lý chức năng đặt lại mật khẩu
- Duy trì trạng thái phiên người dùng""",

    "services/borrow_service.dart": """MỤC ĐÍCH FILE:
Lớp service để quản lý các thao tác mượn thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Tạo và cập nhật yêu cầu mượn
- Quản lý quy trình phê duyệt/từ chối
- Xử lý quy trình trả thiết bị
- Truy xuất lịch sử mượn và yêu cầu đang hoạt động""",

    "services/data_service.dart": """MỤC ĐÍCH FILE:
Service dữ liệu cốt lõi cung cấp các thao tác CRUD cho tất cả các thực thể database.

CÁC CHỨC NĂNG CHÍNH:
- Lớp truy cập dữ liệu chung cho các thao tác Supabase
- Xử lý các thao tác tạo, đọc, cập nhật, xóa
- Cung cấp truy xuất dữ liệu với bộ lọc và sắp xếp
- Quản lý mối quan hệ giữa các thực thể""",

    "services/email_notification_service.dart": """MỤC ĐÍCH FILE:
Quản lý thông báo email cho các sự kiện hệ thống khác nhau.

CÁC CHỨC NĂNG CHÍNH:
- Gửi thông báo email cho yêu cầu mượn
- Thông báo cho người dùng về các sự kiện phê duyệt/từ chối/trả
- Tích hợp với nhà cung cấp dịch vụ email
- Xử lý mẫu và định dạng email""",

    "services/equipment_identifier_service.dart": """MỤC ĐÍCH FILE:
Service để tạo và xác thực mã định danh thiết bị duy nhất.

CÁC CHỨC NĂNG CHÍNH:
- Tạo ID thiết bị duy nhất
- Xác thực định dạng mã định danh
- Đảm bảo tính duy nhất trong toàn hệ thống
- Quản lý các schema mã định danh""",

    "services/excel_import_service.dart": """MỤC ĐÍCH FILE:
Xử lý nhập dữ liệu thiết bị từ file Excel.

CÁC CHỨC NĂNG CHÍNH:
- Phân tích file Excel chứa dữ liệu thiết bị
- Xác thực dữ liệu được nhập
- Nhập hàng loạt thiết bị vào database
- Cung cấp xem trước nhập và báo cáo lỗi""",

    "services/metadata_service.dart": """MỤC ĐÍCH FILE:
Quản lý metadata hệ thống và dữ liệu tham chiếu.

CÁC CHỨC NĂNG CHÍNH:
- Truy xuất dữ liệu tra cứu (danh mục, trạng thái, v.v.)
- Cache metadata được truy cập thường xuyên
- Cung cấp tùy chọn dropdown cho các form""",

    "services/qr_code_service.dart": """MỤC ĐÍCH FILE:
Xử lý tạo mã QR và các thao tác quét cho thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Tạo mã QR cho thiết bị
- Quét và giải mã mã QR
- Liên kết mã QR với bản ghi thiết bị
- Cho phép tra cứu thiết bị nhanh qua quét QR""",

    "services/statistics_service.dart": """MỤC ĐÍCH FILE:
Cung cấp dữ liệu thống kê và phân tích cho dashboard và báo cáo.

CÁC CHỨC NĂNG CHÍNH:
- Tính toán thống kê sử dụng thiết bị
- Tạo xu hướng mượn và phân tích
- Cung cấp dữ liệu cho biểu đồ và báo cáo
- Tổng hợp các chỉ số toàn hệ thống""",

    "services/user_service.dart": """MỤC ĐÍCH FILE:
Lớp service để quản lý tài khoản và hồ sơ người dùng.

CÁC CHỨC NĂNG CHÍNH:
- Tạo và cập nhật hồ sơ người dùng
- Quản lý vai trò và quyền hạn của người dùng
- Truy xuất thông tin người dùng
- Xử lý các thao tác tài khoản người dùng""",

    "services/user_settings_service.dart": """MỤC ĐÍCH FILE:
Quản lý tùy chọn người dùng và cài đặt ứng dụng.

CÁC CHỨC NĂNG CHÍNH:
- Lưu và truy xuất cài đặt người dùng
- Quản lý tùy chọn ngôn ngữ
- Xử lý cài đặt thông báo
- Lưu trữ các tùy chỉnh của người dùng""",

    # Utils
    "utils/equipment_identifiers.dart": """MỤC ĐÍCH FILE:
Các hàm tiện ích để xử lý các thao tác mã định danh thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Xác thực định dạng mã định danh
- Phân tích và định dạng ID
- Các hàm tiện ích mã định danh""",

    "utils/equipment_utils.dart": """MỤC ĐÍCH FILE:
Các hàm tiện ích chung cho các thao tác liên quan đến thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Các thao tác thiết bị phổ biến
- Các hàm helper để thao tác dữ liệu thiết bị
- Các tiện ích định dạng cụ thể cho thiết bị""",

    "utils/equipment_validation.dart": """MỤC ĐÍCH FILE:
Logic xác thực cho dữ liệu thiết bị và form.

CÁC CHỨC NĂNG CHÍNH:
- Xác thực dữ liệu đầu vào thiết bị
- Kiểm tra các trường bắt buộc và định dạng
- Cung cấp thông báo lỗi xác thực
- Đảm bảo tính toàn vẹn dữ liệu""",

    "utils/logger.dart": """MỤC ĐÍCH FILE:
Tiện ích ghi log để gỡ lỗi và giám sát hành vi ứng dụng.

CÁC CHỨC NĂNG CHÍNH:
- Cung cấp ghi log có cấu trúc
- Ghi log lỗi, cảnh báo và thông tin
- Hỗ trợ gỡ lỗi và khắc phục sự cố""",

    "utils/serial_generator.dart": """MỤC ĐÍCH FILE:
Tạo số serial duy nhất cho thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Tạo số serial duy nhất
- Đảm bảo tính duy nhất của số serial
- Tuân theo định dạng số serial có thể cấu hình""",

    # Widgets
    "widgets/continuous_scan_popup.dart": """MỤC ĐÍCH FILE:
Widget UI cho chức năng quét mã QR liên tục.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị máy quét QR trong dialog popup
- Cho phép quét liên tục mà không cần đóng
- Cung cấp phản hồi cho các lần quét thành công""",

    "widgets/equipment_card.dart": """MỤC ĐÍCH FILE:
Widget card có thể tái sử dụng để hiển thị thông tin thiết bị trong danh sách.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị tóm tắt thiết bị (tên, trạng thái, vị trí)
- Cung cấp layout card nhất quán
- Xử lý tương tác chạm để xem chi tiết thiết bị""",

    "widgets/error_dialog.dart": """MỤC ĐÍCH FILE:
Widget dialog có thể tái sử dụng để hiển thị thông báo lỗi.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị thông báo lỗi thân thiện với người dùng
- Cung cấp UI lỗi nhất quán
- Xử lý việc đóng lỗi""",

    "widgets/grouped_borrow_request_card.dart": """MỤC ĐÍCH FILE:
Hiển thị nhiều yêu cầu mượn được nhóm lại với nhau trong định dạng card.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị thông tin yêu cầu mượn được nhóm
- Cung cấp chế độ xem có thể mở rộng/thu gọn
- Xử lý các thao tác hàng loạt trên các yêu cầu được nhóm""",

    "widgets/loading_indicator.dart": """MỤC ĐÍCH FILE:
Widget chỉ báo tải có thể tái sử dụng cho các thao tác bất đồng bộ.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị spinner tải trong khi truy xuất dữ liệu
- Cung cấp UI tải nhất quán
- Thông báo tải có thể tùy chỉnh""",

    "widgets/qr_scanner_widget.dart": """MỤC ĐÍCH FILE:
Component widget máy quét mã QR cốt lõi.

CÁC CHỨC NĂNG CHÍNH:
- Tích hợp camera để quét QR
- Xử lý phát hiện và phân tích mã QR
- Cung cấp callback kết quả quét""",

    "widgets/qr_scan_return_dialog.dart": """MỤC ĐÍCH FILE:
Widget dialog để trả thiết bị qua quét mã QR.

CÁC CHỨC NĂNG CHÍNH:
- Kết hợp quét QR với quy trình trả
- Xác thực thiết bị được quét để trả
- Xác nhận các thao tác trả""",

    # Screens - Admin
    "screens/admin/admin_dashboard_screen.dart": """MỤC ĐÍCH FILE:
Màn hình dashboard chính cho quản trị viên hiển thị tổng quan hệ thống và các hành động nhanh.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị các chỉ số và thống kê dành riêng cho admin
- Cung cấp điều hướng đến các tính năng admin
- Hiển thị tình trạng hệ thống và cảnh báo""",

    "screens/admin/analytics_screen.dart": """MỤC ĐÍCH FILE:
Màn hình phân tích và báo cáo hiển thị biểu đồ và thống kê.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị phân tích sử dụng và xu hướng
- Hiển thị biểu đồ và đồ thị
- Cung cấp chức năng xuất dữ liệu
- Lọc dữ liệu theo khoảng thời gian""",

    "screens/admin/audit_logs_screen.dart": """MỤC ĐÍCH FILE:
Màn hình để xem và tìm kiếm nhật ký kiểm toán hệ thống.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị danh sách theo thời gian các hành động hệ thống
- Cung cấp chức năng tìm kiếm và lọc
- Hiển thị thông tin kiểm toán chi tiết
- Hỗ trợ xuất nhật ký kiểm toán""",

    "screens/admin/category_management_screen.dart": """MỤC ĐÍCH FILE:
Màn hình để quản lý danh mục thiết bị (các thao tác CRUD).

CÁC CHỨC NĂNG CHÍNH:
- Liệt kê tất cả danh mục thiết bị
- Tạo danh mục mới
- Chỉnh sửa danh mục hiện có
- Xóa danh mục không sử dụng""",

    "screens/admin/user_management_screen.dart": """MỤC ĐÍCH FILE:
Màn hình để quản lý tài khoản người dùng và quyền hạn.

CÁC CHỨC NĂNG CHÍNH:
- Liệt kê tất cả người dùng hệ thống
- Tạo tài khoản người dùng mới
- Chỉnh sửa vai trò và quyền hạn người dùng
- Vô hiệu hóa/kích hoạt tài khoản người dùng""",

    # Screens - Auth
    "screens/auth/sign_in_screen.dart": """MỤC ĐÍCH FILE:
Màn hình xác thực người dùng để đăng nhập vào ứng dụng.

CÁC CHỨC NĂNG CHÍNH:
- Cung cấp form đăng nhập bằng email/mật khẩu
- Xử lý gửi xác thực
- Hiển thị thông báo lỗi cho đăng nhập thất bại
- Cung cấp liên kết đến đăng ký và đặt lại mật khẩu
- Bao gồm tùy chọn truy cập khách""",

    "screens/auth/sign_up_screen.dart": """MỤC ĐÍCH FILE:
Màn hình đăng ký người dùng để tạo tài khoản mới.

CÁC CHỨC NĂNG CHÍNH:
- Cung cấp form đăng ký với các trường bắt buộc
- Xác thực đầu vào người dùng
- Tạo tài khoản người dùng mới
- Xử lý lỗi đăng ký
- Chuyển hướng đến đăng nhập sau khi đăng ký thành công""",

    # Screens - Borrow
    "screens/borrow/borrow_list_tab.dart": """MỤC ĐÍCH FILE:
Tab hiển thị các yêu cầu mượn đang hoạt động trong màn hình quản lý mượn.

CÁC CHỨC NĂNG CHÍNH:
- Liệt kê các yêu cầu mượn đang chờ và đã phê duyệt
- Cung cấp tùy chọn lọc và sắp xếp
- Hiển thị chi tiết và trạng thái yêu cầu
- Cho phép các hành động phê duyệt/từ chối yêu cầu""",

    "screens/borrow/borrow_management_screen.dart": """MỤC ĐÍCH FILE:
Màn hình chính để quản lý tất cả yêu cầu mượn với các tab cho các chế độ xem khác nhau.

CÁC CHỨC NĂNG CHÍNH:
- Container cho các tab yêu cầu mượn
- Cung cấp điều hướng giữa yêu cầu đang hoạt động và đã trả
- Hiển thị thống kê tóm tắt
- Cho phép các hành động hàng loạt trên yêu cầu""",

    "screens/borrow/returned_requests_tab.dart": """MỤC ĐÍCH FILE:
Tab hiển thị các yêu cầu mượn đã trả/hoàn thành.

CÁC CHỨC NĂNG CHÍNH:
- Liệt kê các bản ghi mượn lịch sử
- Cung cấp lọc theo ngày và người dùng
- Hiển thị chi tiết và thời gian trả
- Cho phép xem chi tiết yêu cầu đã hoàn thành""",

    "screens/borrow/return_equipment_dialog.dart": """MỤC ĐÍCH FILE:
Dialog để xử lý việc trả thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị form trả thiết bị
- Xác thực điều kiện trả
- Ghi lại thời gian và ghi chú trả
- Cập nhật trạng thái thiết bị thành có sẵn""",

    # Screens - Dashboard
    "screens/dashboard/main_dashboard.dart": """MỤC ĐÍCH FILE:
Màn hình dashboard chính với ngăn kéo điều hướng và các tùy chọn menu dựa trên vai trò.

CÁC CHỨC NĂNG CHÍNH:
- Cung cấp cấu trúc điều hướng chính
- Hiển thị các tùy chọn khác nhau dựa trên vai trò người dùng
- Hiển thị các chỉ số và thống kê chính
- Cho phép truy cập nhanh vào các chức năng phổ biến""",

    # Screens - Equipment
    "screens/equipment/equipment_catalog_screen.dart": """MỤC ĐÍCH FILE:
Màn hình hiển thị danh mục có thể tìm kiếm của tất cả thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Liệt kê tất cả thiết bị với tìm kiếm/lọc
- Hiển thị trạng thái sẵn có của thiết bị
- Cho phép quét mã QR để tra cứu nhanh
- Cung cấp điều hướng đến chi tiết thiết bị
- Cho phép tạo yêu cầu mượn""",

    "screens/equipment/equipment_form_screen.dart": """MỤC ĐÍCH FILE:
Màn hình form để tạo và chỉnh sửa bản ghi thiết bị.

CÁC CHỨC NĂNG CHÍNH:
- Cung cấp form nhập cho chi tiết thiết bị
- Xác thực dữ liệu thiết bị
- Tạo mã định danh thiết bị
- Lưu thiết bị mới hoặc đã cập nhật
- Tạo mã QR cho thiết bị""",

    "screens/equipment/equipment_import_preview_screen.dart": """MỤC ĐÍCH FILE:
Màn hình xem trước để nhập hàng loạt thiết bị từ file Excel.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị dữ liệu Excel đã phân tích
- Hiển thị lỗi và cảnh báo xác thực
- Cho phép chỉnh sửa trước khi nhập
- Thực hiện nhập hàng loạt vào database
- Báo cáo thống kê nhập thành công/thất bại""",

    # Screens - Settings
    "screens/settings/settings_screen.dart": """MỤC ĐÍCH FILE:
Màn hình cài đặt người dùng để quản lý tùy chọn và hồ sơ.

CÁC CHỨC NĂNG CHÍNH:
- Hiển thị thông tin hồ sơ người dùng
- Cho phép lựa chọn ngôn ngữ
- Quản lý tùy chọn thông báo
- Cung cấp chức năng đăng xuất
- Hiển thị phiên bản ứng dụng và thông tin về""",
}

def create_word_document(file_path, file_key, explanation, output_dir):
    """Create a Word document for a single Dart file"""
    try:
        # Read file content
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create Word document
        doc = Document()
        
        # File name
        file_heading = doc.add_heading(f'File: {file_key}', level=1)
        file_run = file_heading.runs[0]
        file_run.font.name = 'Times New Roman'
        file_run.font.size = Pt(13)
        file_run.font.color.rgb = RGBColor(0, 0, 255)
        
        doc.add_paragraph()
        
        # Explanation section
        explanation_heading = doc.add_heading('GIẢI THÍCH FILE:', level=2)
        explanation_run = explanation_heading.runs[0]
        explanation_run.font.name = 'Times New Roman'
        explanation_run.font.size = Pt(13)
        explanation_run.font.color.rgb = RGBColor(0, 100, 200)
        
        explanation_para = doc.add_paragraph(explanation)
        explanation_para.style.font.name = 'Times New Roman'
        explanation_para.style.font.size = Pt(13)
        explanation_para.paragraph_format.line_spacing = 1.15
        for run in explanation_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        
        doc.add_paragraph()
        
        # Code section
        code_heading = doc.add_heading('MÃ NGUỒN:', level=2)
        code_run = code_heading.runs[0]
        code_run.font.name = 'Times New Roman'
        code_run.font.size = Pt(13)
        code_run.font.color.rgb = RGBColor(0, 100, 200)
        
        code_para = doc.add_paragraph(content)
        code_para.style.font.name = 'Times New Roman'
        code_para.style.font.size = Pt(13)
        code_para.paragraph_format.line_spacing = 1.15
        for run in code_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        
        # Save document
        output_filename = file_key.replace('/', '_').replace('\\', '_').replace('.dart', '.docx')
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        print(f"✓ Đã tạo: {output_filename}")
        return True
        
    except Exception as e:
        print(f"✗ Lỗi khi xử lý {file_key}: {str(e)}")
        return False

def main():
    lib_path = r"c:\Users\PC\Documents\DATN\medical\flutter_application_1\lib"
    output_path = r"c:\Users\PC\Documents\DATN\medical\flutter_application_1\lib_documentation"
    
    # Create output directory
    os.makedirs(output_path, exist_ok=True)
    
    print("=" * 60)
    print("Bắt đầu tạo tài liệu Word...")
    print("=" * 60)
    print()
    
    success_count = 0
    fail_count = 0
    
    for file_key, explanation in file_explanations.items():
        file_path = os.path.join(lib_path, file_key)
        
        if os.path.exists(file_path):
            if create_word_document(file_path, file_key, explanation, output_path):
                success_count += 1
            else:
                fail_count += 1
        else:
            print(f"⚠ Không tìm thấy file: {file_key}")
            fail_count += 1
    
    print()
    print("=" * 60)
    print(f"✓ Đã tạo thành công: {success_count} tài liệu")
    if fail_count > 0:
        print(f"✗ Thất bại: {fail_count} tài liệu")
    print(f"📁 Vị trí lưu: {output_path}")
    print("=" * 60)

if __name__ == "__main__":
    main()
